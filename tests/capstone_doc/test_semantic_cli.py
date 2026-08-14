from __future__ import annotations

import json
from pathlib import Path

from docx import Document as WordDocument

from rift_doc.cli import main
from rift_doc.engine import ValidationEngine
from rift_doc.semantic import FakeLLMProvider, SemanticAuditOptions
from rift_doc.spec import CapstoneSpec


_SPEC = """\
$schema: capstone-doc-spec.schema.json
spec_version: 'semantic-cli-test-1'
source:
  authority: test
  documents: []
classification:
  placeholder_patterns: []
  instruction_patterns: []
  sample_fingerprints: []
reports:
  report3:
    title: Requirements
    source_requirement: Functions describe normal and abnormal behavior.
    sections:
      feature_requirements:
        title: Feature Requirements
        requirement: MUST
        match:
          regex: '^3\\.\\d+\\s+.+$'
        repeatable: true
        content: true
workbooks: {}
semantic_review_extension:
  source_precedence:
    authoritative_sources: [report3]
  rules:
    - id: SEM-R3-CONTENT
      task_type: CONTENT_SUFFICIENCY
      generation: sections
      source_domain: report3
      section_rules: [feature_requirements]
      question: 'Does {source_name} cover normal and abnormal behavior?'
      prompt_version: content_sufficiency.v1
      allowed_statuses: [PASS, FAIL, REVIEW_REQUIRED]
source_ambiguities: []
"""


def _audit_files(tmp_path: Path) -> tuple[Path, Path]:
    spec_path = tmp_path / "spec.yaml"
    spec_path.write_text(_SPEC, encoding="utf-8")
    document_path = tmp_path / "report3.docx"
    document = WordDocument()
    document.add_heading("3.2 Notification Sync", level=1)
    document.add_paragraph(
        "A trusted peer forwards a notification. Invalid payloads are rejected and the user is notified."
    )
    document.save(document_path)
    manifest_path = tmp_path / "manifest.yaml"
    manifest_path.write_text(
        "project: Rift\nreports:\n  report3: report3.docx\n",
        encoding="utf-8",
    )
    return spec_path, manifest_path


def _provider_response(task: object, packet: object) -> dict[str, object]:
    cited = packet.document_evidence[0].evidence_id
    return {
        "status": "PASS",
        "confidence": "HIGH",
        "summary": "The function describes both successful and invalid-payload behavior.",
        "reasoning_summary": "The exact section states the successful action and the rejection outcome.",
        "evidence_refs": [cited],
        "unsupported_claims": [],
        "contradictions": [],
        "recommended_action": "Keep the behavior synchronized with tests and design.",
    }


def test_semantic_plan_cli_is_offline_and_lists_bounded_packets(
    tmp_path: Path,
    capsys: object,
) -> None:
    spec, manifest = _audit_files(tmp_path)

    code = main(
        [
            "semantic-plan",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--task",
            "content-sufficiency",
            "--max-tasks",
            "2",
            "--format",
            "json",
        ]
    )

    assert code == 0
    payload = json.loads(capsys.readouterr().out)
    assert payload["selected_task_count"] == 1
    assert payload["metadata"]["provider_called"] is False
    assert payload["metadata"]["network_access"] is False
    assert payload["packets"][0]["budget"]["estimated_input_tokens"] <= 12_000
    assert payload["packets"][0]["document_evidence"][0]["source_path"].endswith("report3.docx")


def test_semantic_cli_appends_domain_separated_finding(
    tmp_path: Path,
    capsys: object,
    monkeypatch: object,
) -> None:
    spec, manifest = _audit_files(tmp_path)
    monkeypatch.setattr(
        "rift_doc.semantic.providers.OpenAICompatibleProvider.review",
        lambda self, task, packet: _provider_response(task, packet),
    )

    code = main(
        [
            "semantic",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--task",
            "content-sufficiency",
            "--semantic-model",
            "local-review",
            "--semantic-endpoint",
            "http://localhost/v1/chat/completions",
            "--semantic-local-only",
            "--no-cache",
            "--format",
            "json",
        ]
    )

    assert code == 0
    payload = json.loads(capsys.readouterr().out)
    validators = [finding["validator"] for finding in payload["findings"]]
    assert "structural" in validators
    assert validators[-1] == "semantic_review"
    assert payload["metadata"]["semantic_review"]["metadata"]["domain"] == "semantic_review"
    assert payload["metadata"]["semantic_review"]["plan"]["metadata"]["dry_run"] is False


def test_validate_set_semantic_text_keeps_deterministic_output_first(
    tmp_path: Path,
    capsys: object,
    monkeypatch: object,
) -> None:
    spec, manifest = _audit_files(tmp_path)
    monkeypatch.setattr(
        "rift_doc.semantic.providers.OpenAICompatibleProvider.review",
        lambda self, task, packet: _provider_response(task, packet),
    )

    code = main(
        [
            "validate-set",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--semantic",
            "--semantic-rule",
            "content-sufficiency",
            "--semantic-model",
            "local-review",
            "--semantic-endpoint",
            "http://localhost/v1/chat/completions",
            "--semantic-local-only",
            "--no-cache",
        ]
    )

    assert code == 0
    output = capsys.readouterr().out
    assert "SEMANTIC REVIEW" in output
    assert "Bounded model review; deterministic findings above remain unchanged." in output
    assert output.index("PASS    report3") < output.index("SEMANTIC REVIEW")
    assert "Evidence:" in output


def test_engine_preserves_deterministic_finding_payloads_exactly(tmp_path: Path) -> None:
    spec_path, manifest = _audit_files(tmp_path)
    engine = ValidationEngine(CapstoneSpec.load(spec_path))
    deterministic = engine.validate_set(manifest)
    original_findings = [finding.to_dict() for finding in deterministic.findings]
    provider = FakeLLMProvider(_provider_response)

    combined = engine.validate_set(
        manifest,
        semantic_provider=provider,
        semantic_options=SemanticAuditOptions(cache_enabled=False),
    )

    assert [finding.to_dict() for finding in combined.findings[: len(original_findings)]] == original_findings
    assert combined.findings[-1].validator == "semantic_review"
    assert combined.metadata["semantic_review"]["deterministic_finding_count"] == len(original_findings)


def test_validate_set_without_semantic_never_constructs_provider(
    tmp_path: Path,
    capsys: object,
    monkeypatch: object,
) -> None:
    spec, manifest = _audit_files(tmp_path)

    def fail_if_constructed(*args: object, **kwargs: object) -> None:
        raise AssertionError("offline deterministic validation constructed a provider")

    monkeypatch.setattr(
        "rift_doc.semantic.providers.OpenAICompatibleProvider.__init__",
        fail_if_constructed,
    )

    code = main(
        [
            "validate-set",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--format",
            "json",
        ]
    )

    assert code == 0
    payload = json.loads(capsys.readouterr().out)
    assert "semantic_review" not in payload["metadata"]


def test_semantic_plan_honors_zero_task_limit(tmp_path: Path, capsys: object) -> None:
    spec, manifest = _audit_files(tmp_path)

    code = main(
        [
            "semantic-plan",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--max-tasks",
            "0",
            "--format",
            "json",
        ]
    )

    assert code == 0
    payload = json.loads(capsys.readouterr().out)
    assert payload["selected_task_count"] == 0
    assert payload["proposed_task_count"] == 1
    assert payload["omitted_task_count"] == 1
