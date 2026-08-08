from __future__ import annotations

from pathlib import Path

import openpyxl
import pytest

from rift_doc.classification import ContentClassifier
from rift_doc.extractors.docx import extract_docx
from rift_doc.extractors.spreadsheets import extract_workbook

from .conftest import SOURCE_TEMPLATES, create_docx


def test_docx_extracts_ordered_sections_tables_images_and_styles(tmp_path: Path) -> None:
    path = create_docx(tmp_path / "fixture.docx", required_text="Completed section content", include_image=True, table=True)
    document = extract_docx(path, classifier=ContentClassifier())
    assert [section.title for section in document.all_sections()] == [
        "Required Section",
        "Recommended Section",
        "Conditional Section",
        "Image Section",
    ]
    assert document.raw_blocks[0].kind == "heading"
    assert document.raw_blocks[1].classification.value == "real_content"
    assert document.tables[0].dimensions == (2, 2)
    assert document.tables[0].parent_section == "Image Section"
    assert len(document.images) == 1
    assert document.images[0].parent_section == "Image Section"
    assert any(block.kind == "image" for block in document.raw_blocks)
    assert document.images[0].source_location.paragraph_index is not None


def test_docx_heading_style_is_used_before_visible_text(tmp_path: Path) -> None:
    from docx import Document as WordDocument

    path = tmp_path / "rift-doc-heading-test.docx"
    document = WordDocument()
    document.add_paragraph("A visible heading without a number", style="Heading 2")
    document.save(path)
    normalized = extract_docx(path)
    assert normalized.sections[0].title == "A visible heading without a number"
    assert normalized.sections[0].level == 2


def test_xlsx_extracts_sheets_cells_merged_regions_headers_and_formulas(tmp_path: Path) -> None:
    path = tmp_path / "fixture.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Sheet A"
    sheet.merge_cells("A1:B1")
    sheet["A1"] = "Title"
    sheet["A2"] = "Name"
    sheet["B2"] = "Status"
    sheet["A3"] = "Real"
    sheet["B3"] = "Pending"
    sheet["C3"] = "=1+1"
    workbook.save(path)

    normalized = extract_workbook(path)
    assert normalized.format == "xlsx"
    assert normalized.sheets[0].name == "Sheet A"
    assert normalized.sheets[0].merged_regions == ["A1:B1"]
    assert normalized.sheets[0].detected_header_rows == [2]
    formula_cell = normalized.sheets[0].rows[2][2]
    assert formula_cell.formula == "=1+1"
    assert formula_cell.value is None  # no cached value; the formula is retained separately
    assert not formula_cell.is_empty


def test_real_report5_xls_is_extractable_without_rewriting_source() -> None:
    path = SOURCE_TEMPLATES / "Report5_Unit Test.xls"
    if not path.exists():
        pytest.skip("source templates were not supplied in this checkout")
    before = path.stat().st_mtime_ns
    normalized = extract_workbook(path)
    after = path.stat().st_mtime_ns
    assert before == after
    assert normalized.format == "xls"
    assert [sheet.name for sheet in normalized.sheets] == [
        "Guideline",
        "Cover",
        "Functions",
        "Statistics",
        "Function 1",
        "Function 2",
        "Function3",
        "Example",
    ]
    assert "B2:F2" in normalized.sheets[1].merged_regions
    assert normalized.sheets[2].detected_header_rows == [10]
    assert normalized.sheets[3].detected_header_rows == [11]


def test_real_report5_xlsx_schema_is_extractable() -> None:
    path = SOURCE_TEMPLATES / "Report5_Test Report.xlsx"
    if not path.exists():
        pytest.skip("source templates were not supplied in this checkout")
    normalized = extract_workbook(path)
    assert [sheet.name for sheet in normalized.sheets] == [
        "Cover",
        "Test Cases",
        "Test Statistics",
        "Feature 1",
        "Feature 2",
    ]
    assert normalized.sheets[1].detected_header_rows == [8]
    assert normalized.sheets[2].detected_header_rows == [10]
    assert "B2:E2" in normalized.sheets[3].merged_regions
