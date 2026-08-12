from __future__ import annotations

import json
from pathlib import Path

from docx import Document as WordDocument

from rift_doc.cli import main


_SPEC = """\
$schema: capstone-doc-spec.schema.json
spec_version: 'test-repository'
source:
  authority: test
  documents: []
classification:
  placeholder_patterns: []
  instruction_patterns: []
  sample_fingerprints: []
reports:
  demo:
    title: Demo
    trace_entities:
      feature:
        sections: [Required Section]
    sections:
      required:
        title: Required Section
        requirement: MUST
        content: true
workbooks: {}
repository_evidence_extension:
  finding_domain: repository_evidence
  claims:
    functions:
      claim_kind: FUNCTION_OR_FEATURE
      source_kinds: [feature]
      source_domains: [demo]
      expected_evidence_types: [SYMBOL, CONFIGURATION]
source_ambiguities: []
"""


def _fixture(tmp_path: Path, *, symbol: str = "notification_sync") -> tuple[Path, Path, Path]:
    spec = tmp_path / "spec.yaml"
    spec.write_text(_SPEC, encoding="utf-8")
    document_path = tmp_path / "report.docx"
    document = WordDocument()
    document.add_heading("Required Section", level=1)
    document.add_paragraph("FE-03 Notification Sync")
    document.save(document_path)
    manifest = tmp_path / "capstone.yaml"
    manifest.write_text("reports:\n  demo: report.docx\n", encoding="utf-8")
    repository = tmp_path / "repo"
    (repository / "lib").mkdir(parents=True)
    (repository / "lib" / "sync.py").write_text(f"def {symbol}():\n    return True\n", encoding="utf-8")
    return spec, manifest, repository


def test_repo_inspect_supports_human_and_json_output(tmp_path: Path, capsys) -> None:
    _spec, _manifest, repository = _fixture(tmp_path)

    assert main(["repo-inspect", "--repo", str(repository)]) == 0
    text = capsys.readouterr().out
    assert "Detected repository" in text
    assert "Commit: unavailable (plain source tree)" in text
    assert "Languages: python (1)" in text
    assert "Symbols: 1" in text

    assert main(["repo-inspect", "--repo", str(repository), "--format", "json"]) == 0
    payload = json.loads(capsys.readouterr().out)
    assert payload["root"] == str(repository.resolve())
    assert payload["vcs_metadata"] is None
    assert payload["symbols"][0]["path"] == "lib/sync.py"
    assert payload["metadata"]["repository_code_executed"] is False


def test_evidence_command_emits_doc_and_repo_locations_in_json(tmp_path: Path, capsys) -> None:
    spec, manifest, repository = _fixture(tmp_path)

    assert main(
        [
            "evidence",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--repo",
            str(repository),
            "--format",
            "json",
        ]
    ) == 0
    payload = json.loads(capsys.readouterr().out)

    finding = next(item for item in payload["findings"] if item["rule_id"] == "REPO-FUNCTION")
    assert finding["validator"] == "repository_evidence"
    assert finding["metadata"]["evidence_status"] == "VERIFIED"
    assert finding["source_entity"]["documentation_evidence"][0]["source_path"] == str(tmp_path / "report.docx")
    assert finding["candidate_entities"][0]["path"] == "lib/sync.py"
    assert finding["metadata"]["repository_locations"] == ["lib/sync.py:1"]
    assert payload["metadata"]["repository_snapshot"]["root"] == str(repository.resolve())
    assert payload["metadata"]["evidence_packets"][0]["deterministic_finding"]["status"] == "VERIFIED"


def test_validate_set_repo_integration_has_separate_human_section(tmp_path: Path, capsys) -> None:
    spec, manifest, repository = _fixture(tmp_path)

    assert main(
        [
            "validate-set",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--repo",
            str(repository),
        ]
    ) == 0
    text = capsys.readouterr().out
    assert "REPOSITORY EVIDENCE" in text
    assert "VERIFIED" in text
    assert "documentation: demo @ paragraph 1" in text
    assert "repository: lib/sync.py:1" in text
    assert "not an FPT template" in text


def test_evidence_filters_claims_and_kinds(tmp_path: Path, capsys) -> None:
    spec, manifest, repository = _fixture(tmp_path)

    assert main(
        [
            "evidence",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--repo",
            str(repository),
            "--kind",
            "architecture",
            "--format",
            "json",
        ]
    ) == 0
    by_kind = json.loads(capsys.readouterr().out)
    assert by_kind["metadata"]["claim_count"] == 0
    assert by_kind["findings"] == []

    assert main(
        [
            "evidence",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--repo",
            str(repository),
            "--claim",
            "FE-03",
            "--format",
            "json",
        ]
    ) == 0
    by_claim = json.loads(capsys.readouterr().out)
    assert by_claim["metadata"]["claim_count"] == 1


def test_cli_manual_mapping_is_auditable_and_invalid_mapping_fails(tmp_path: Path, capsys) -> None:
    spec, manifest, repository = _fixture(tmp_path, symbol="perform_remote_sync")
    mapping = tmp_path / "mapping.yaml"
    mapping.write_text(
        """
repository_mappings:
  features:
    FE-03:
      symbols: [perform_remote_sync]
""",
        encoding="utf-8",
    )

    assert main(
        [
            "evidence",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--repo",
            str(repository),
            "--mapping",
            str(mapping),
            "--format",
            "json",
        ]
    ) == 0
    payload = json.loads(capsys.readouterr().out)
    finding = next(item for item in payload["findings"] if item["rule_id"] == "REPO-FUNCTION")
    assert finding["metadata"]["match_method"] == "MANUAL_MAPPING"
    assert finding["metadata"]["manual_mapping"] is True
    assert finding["metadata"]["mapping"]["source_path"] == str(mapping)

    mapping.write_text("repository_mappings:\n  features:\n    FE-03:\n      unknown: true\n", encoding="utf-8")
    assert main(
        [
            "evidence",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--repo",
            str(repository),
            "--mapping",
            str(mapping),
        ]
    ) == 2
    assert "invalid repository mapping" in capsys.readouterr().err


def test_validate_set_without_repo_preserves_phase_one_two_shape(tmp_path: Path, capsys) -> None:
    spec, manifest, _repository = _fixture(tmp_path)

    assert main(
        [
            "validate-set",
            "--spec",
            str(spec),
            "--manifest",
            str(manifest),
            "--format",
            "json",
        ]
    ) == 0
    payload = json.loads(capsys.readouterr().out)
    assert "phase1" in payload["metadata"]
    assert "trace_graph" in payload["metadata"]
    assert "repository_evidence" not in payload["metadata"]
    assert all(item["validator"] != "repository_evidence" for item in payload["findings"])
