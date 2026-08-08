from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as WordDocument

from rift_doc.engine import ValidationEngine
from rift_doc.extractors.docx import extract_docx
from rift_doc.results import Status
from rift_doc.spec import CapstoneSpec

from .conftest import ROOT, create_docx


def _create_report1_fixture(path: Path, *, manual_headings: bool) -> Path:
    document = WordDocument()

    def heading(text: str, level: int) -> None:
        if manual_headings:
            document.add_paragraph(text)
        else:
            document.add_heading(text, level=level)

    def content(text: str) -> None:
        document.add_paragraph(text)

    def table() -> None:
        value = document.add_table(rows=2, cols=2)
        value.cell(0, 0).text = "Header"
        value.cell(0, 1).text = "Value"
        value.cell(1, 0).text = "Project"
        value.cell(1, 1).text = "Completed"

    heading("I. Record of Changes", 1)
    table()
    heading("II. Project Introduction", 1)
    heading("1. Overview", 2)
    heading("1.1 Project Information", 3)
    content("The project information describes the completed project context.")
    heading("1.2 Project Team", 3)
    content("The project team and its responsibilities are documented.")
    table()
    heading("2. Product Background", 2)
    content("The product background explains the project need.")
    heading("3. Existing Systems", 2)
    content("Existing systems and their relevant capabilities are compared.")
    heading("3.1 Existing System", 3)
    content("The first existing system is documented for comparison.")
    heading("3.2 Existing System", 3)
    content("The second existing system is documented for comparison.")
    heading("4. Business Opportunity", 2)
    content("The business opportunity explains the expected value.")
    heading("5. Software Product Vision", 2)
    content("The product vision describes the intended outcome.")
    heading("6. Project Scope & Limitations", 2)
    heading("6.1 Major Features", 3)
    content("The major features define the planned project scope.")
    heading("6.2 Limitations & Exclusions", 3)
    content("The limitations and exclusions define the project boundaries.")
    document.save(path)
    return path


@pytest.mark.parametrize("manual_headings", [False, True], ids=["styled", "manually-numbered"])
def test_report1_sections_are_nested_for_styled_and_manual_headings(
    manual_headings: bool,
    tmp_path: Path,
) -> None:
    path = _create_report1_fixture(tmp_path / "report1.docx", manual_headings=manual_headings)
    spec = CapstoneSpec.load(ROOT / "capstone-doc-spec.v0.1.yaml")
    engine = ValidationEngine(spec)
    document = engine.extract(path, "report1")
    introduction = next(section for section in document.sections if section.normalized_title == "project introduction")
    assert [section.normalized_title for section in introduction.children] == [
        "overview",
        "product background",
        "existing systems",
        "business opportunity",
        "software product vision",
        "project scope & limitations",
    ]
    assert engine.structural.known_headings("report1")["Project Introduction"] == 1
    result = engine.validate_normalized(document, "report1")
    assert not any(item.status == Status.FAIL for item in result.findings)


def _engine(minimal_spec_path: Path) -> ValidationEngine:
    return ValidationEngine(CapstoneSpec.load(minimal_spec_path))


def test_heading_without_real_content_fails_required_content(minimal_spec_path: Path, tmp_path: Path) -> None:
    path = create_docx(tmp_path / "empty.docx")
    result = _engine(minimal_spec_path).validate(path, "demo")
    finding = next(item for item in result.findings if item.rule_id == "required.content")
    assert finding.status == Status.FAIL
    assert "no real completed content" in finding.message


def test_instruction_only_section_does_not_count_as_content(minimal_spec_path: Path, tmp_path: Path) -> None:
    path = create_docx(tmp_path / "instruction.docx", required_text="[Provide the completed section]")
    result = _engine(minimal_spec_path).validate(path, "demo")
    assert any(item.rule_id == "template.instruction" and item.status == Status.FAIL for item in result.findings)
    assert any(item.rule_id == "required.content" and item.status == Status.FAIL for item in result.findings)


def test_placeholder_embedded_in_real_prose_is_preserved_and_invalid(minimal_spec_path: Path, tmp_path: Path) -> None:
    path = create_docx(tmp_path / "placeholder.docx", required_text="The service handles <Feature Name> correctly.")
    result = _engine(minimal_spec_path).validate(path, "demo")
    finding = next(item for item in result.findings if item.rule_id == "placeholder.unresolved")
    assert finding.status == Status.FAIL
    assert "Feature Name" in finding.evidence[0]["text"]


def test_numbered_equivalent_heading_is_accepted(minimal_spec_path: Path, tmp_path: Path) -> None:
    document = WordDocument()
    document.add_heading("1 Required Section", level=1)
    document.add_paragraph("Completed content")
    document.add_heading("Recommended Section", level=1)
    document.add_paragraph("Recommended content")
    document.add_heading("Conditional Section", level=1)
    document.add_paragraph("Applicable content")
    document.add_heading("Image Section", level=1)
    document.add_paragraph("Image explanation")
    path = tmp_path / "numbered.docx"
    document.save(path)
    result = _engine(minimal_spec_path).validate(path, "demo")
    assert not any(item.rule_id == "required" and item.status == Status.FAIL for item in result.findings)
    assert any(item.rule_id == "required.content" and item.status == Status.PASS for item in result.findings)


def test_unmapped_numbered_heading_is_reviewable(minimal_spec_path: Path, tmp_path: Path) -> None:
    document = WordDocument()
    document.add_paragraph("9.1 Unmapped Section")
    path = tmp_path / "unmapped.docx"
    document.save(path)
    result = _engine(minimal_spec_path).validate(path, "demo")
    assert any(item.rule_id == "heading.unresolved" and item.status == Status.REVIEW_REQUIRED for item in result.findings)


def test_missing_must_and_should_have_different_statuses(minimal_spec_path: Path, tmp_path: Path) -> None:
    document = WordDocument()
    document.add_paragraph("Only unrelated prose")
    path = tmp_path / "missing.docx"
    document.save(path)
    result = _engine(minimal_spec_path).validate(path, "demo")
    assert next(item for item in result.findings if item.rule_id == "required").status == Status.FAIL
    should = next(item for item in result.findings if item.rule_id == "should")
    assert should.status == Status.WARNING
    assert should.severity == "warning"


def test_conditional_unknown_is_review_required(minimal_spec_path: Path, tmp_path: Path) -> None:
    document = WordDocument()
    document.add_heading("Required Section", level=1)
    document.add_paragraph("Completed content")
    path = tmp_path / "conditional-unknown.docx"
    document.save(path)
    result = _engine(minimal_spec_path).validate(path, "demo")
    conditional = next(item for item in result.findings if item.rule_id == "conditional")
    assert conditional.status == Status.REVIEW_REQUIRED


def test_conditional_true_missing_section_fails(minimal_spec_path: Path, tmp_path: Path) -> None:
    document = WordDocument()
    document.add_heading("Required Section", level=1)
    document.add_paragraph("Completed content")
    path = tmp_path / "conditional-true.docx"
    document.save(path)
    spec_path = tmp_path / "conditional-true.yaml"
    spec_path.write_text(
        minimal_spec_path.read_text(encoding="utf-8").replace("type: unknown", "type: always_true"),
        encoding="utf-8",
    )
    result = ValidationEngine(CapstoneSpec.load(spec_path)).validate(path, "demo")
    conditional = next(item for item in result.findings if item.rule_id == "conditional")
    assert conditional.status == Status.FAIL


def test_unrelated_explicit_na_does_not_waive_missing_conditional(minimal_spec_path: Path, tmp_path: Path) -> None:
    document = WordDocument()
    document.add_heading("Required Section", level=1)
    document.add_paragraph("Completed content")
    document.add_paragraph("The conditional section is not applicable because this product has no such component.")
    path = tmp_path / "conditional-na.docx"
    document.save(path)
    result = _engine(minimal_spec_path).validate(path, "demo")
    conditional = next(item for item in result.findings if item.rule_id == "conditional")
    assert conditional.status == Status.REVIEW_REQUIRED
    assert conditional.status != Status.NOT_APPLICABLE


def test_image_presence_is_pass_but_semantics_remain_reviewable(minimal_spec_path: Path, tmp_path: Path) -> None:
    path = create_docx(tmp_path / "image.docx", required_text="Completed content", include_image=True)
    result = _engine(minimal_spec_path).validate(path, "demo")
    assert any(item.rule_id == "with_image.evidence.0" and item.status == Status.PASS for item in result.findings)
    assert any(item.rule_id == "with_image.evidence.0.semantic" and item.status == Status.REVIEW_REQUIRED for item in result.findings)


def test_table_with_only_headers_is_not_completed_content(minimal_spec_path: Path, tmp_path: Path) -> None:
    document = WordDocument()
    document.add_heading("Required Section", level=1)
    document.add_paragraph("Content exists")
    document.add_heading("Image Section", level=1)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    path = tmp_path / "header-table.docx"
    document.save(path)
    # The base contract does not require a table for Image Section.  Verify the
    # normalized table helper directly through the section evidence path by
    # adding a small contract rule in a copy of the fixture contract.
    spec_text = minimal_spec_path.read_text(encoding="utf-8").replace(
        "        content: true\n        evidence:\n          - type: image",
        "        content: true\n        evidence:\n          - type: table\n            required: true\n          - type: image",
    )
    spec_path = tmp_path / "table-spec.yaml"
    spec_path.write_text(spec_text, encoding="utf-8")
    result = ValidationEngine(CapstoneSpec.load(spec_path)).validate(path, "demo")
    finding = next(item for item in result.findings if item.rule_id == "with_image.evidence.0")
    assert finding.status == Status.FAIL
