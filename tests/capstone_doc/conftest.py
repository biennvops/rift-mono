from __future__ import annotations

from pathlib import Path
from io import BytesIO
import base64

import pytest
from docx import Document as WordDocument
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
SCHEMA = ROOT / "capstone-doc-spec.schema.json"
SOURCE_TEMPLATES = ROOT / "capstone-documents" / "source-templates"


MINIMAL_SPEC = """\
$schema: capstone-doc-spec.schema.json
spec_version: 'test-1'
source:
  authority: test
  documents: []
classification:
  placeholder_patterns: []
  instruction_patterns: []
  sample_fingerprints: []
reports:
  demo:
    title: Demo
    sections:
      required:
        title: Required Section
        requirement: MUST
        content: true
      should:
        title: Recommended Section
        requirement: SHOULD
        content: true
      conditional:
        title: Conditional Section
        requirement: CONDITIONAL
        condition:
          type: unknown
        allow_explicit_na: true
        content: true
      with_image:
        title: Image Section
        requirement: MUST
        content: true
        evidence:
          - type: image
            required: true
            semantic_review: true
workbooks: {}
source_ambiguities: []
"""


@pytest.fixture
def minimal_spec_path(tmp_path: Path) -> Path:
    path = tmp_path / "demo-spec.yaml"
    path.write_text(MINIMAL_SPEC, encoding="utf-8")
    return path


def create_docx(path: Path, *, required_text: str | None = None, include_image: bool = False, table: bool = False, complete_optional: bool = False) -> Path:
    document = WordDocument()
    document.add_heading("Required Section", level=1)
    if required_text is not None:
        document.add_paragraph(required_text)
    document.add_heading("Recommended Section", level=1)
    if complete_optional:
        document.add_paragraph("Recommended content")
    document.add_heading("Conditional Section", level=1)
    if complete_optional:
        document.add_paragraph("Conditional content applies")
    document.add_heading("Image Section", level=1)
    if include_image:
        document.add_picture(BytesIO(_PNG_BYTES), width=Inches(1))
        document.add_paragraph("Image explanation")
    if table:
        table_object = document.add_table(rows=2, cols=2)
        table_object.cell(0, 0).text = "Header"
        table_object.cell(0, 1).text = "Value"
        table_object.cell(1, 0).text = "Real"
        table_object.cell(1, 1).text = "Content"
    document.save(path)
    return path


# 1x1 transparent PNG; keeping the fixture inline avoids an image dependency.
_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)
