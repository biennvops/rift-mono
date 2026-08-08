"""Deterministic DOCX extraction built on python-docx."""

from __future__ import annotations

from collections.abc import Iterable, Mapping
import re
from pathlib import Path
import unicodedata
from typing import Any

from ..classification import ContentClassifier
from ..model import Block, Document, Image, Section, SourceLocation, Table


_NUMBER_PREFIX = re.compile(
    r"^\s*(?P<number>(?:[IVXLCDM]+\.|\d+(?:\.\d+)*\.?))\s+(?P<title>.+?)\s*$",
    re.IGNORECASE,
)
_MANUAL_HEADING = re.compile(r"^\s*\d+(?:\.\d+)+\s+\S.{0,180}$")


def normalize_heading(title: str) -> str:
    """Normalize numbering and trivial typography, without fuzzy matching."""

    value = unicodedata.normalize("NFKC", title or "")
    value = value.replace("\u00a0", " ").replace("\u2013", "-").replace("\u2014", "-")
    value = re.sub(r"\s+", " ", value).strip()
    match = _NUMBER_PREFIX.match(value)
    if match:
        value = match.group("title")
    value = value.strip(" .:-\t")
    value = re.sub(r"[.!:]+$", "", value).strip()
    return value.casefold()


def extract_numbering(text: str) -> str | None:
    match = _NUMBER_PREFIX.match(text or "")
    return match.group("number") if match else None


def _numbering_level(numbering: str | None) -> int | None:
    if not numbering:
        return None
    if re.fullmatch(r"[IVXLCDM]+\.", numbering, re.IGNORECASE):
        return 1
    return len([part for part in numbering.rstrip(".").split(".") if part]) + 1


class HeadingMatcher:
    """Exact normalized heading matching supplied by the contract."""

    def __init__(self, known_headings: Mapping[str, int] | Iterable[str] | None = None) -> None:
        self.known: dict[str, int] = {}
        if isinstance(known_headings, Mapping):
            for title, level in known_headings.items():
                self.known[normalize_heading(str(title))] = int(level)
        else:
            for title in known_headings or ():
                self.known[normalize_heading(str(title))] = _numbering_level(extract_numbering(str(title))) or 2

    def match(self, text: str) -> int | None:
        normalized = normalize_heading(text)
        if normalized in self.known:
            return self.known[normalized]
        # This is deliberately limited to a clearly numbered, short heading.
        # It does not attempt synonym or semantic matching.
        if _MANUAL_HEADING.match(text) and not text.rstrip().endswith((".", ":", "]")):
            return _numbering_level(extract_numbering(text)) or 2
        return None


def extract_docx(
    path: str | Path,
    *,
    classifier: ContentClassifier | None = None,
    known_headings: Mapping[str, int] | Iterable[str] | None = None,
) -> Document:
    """Extract body blocks, sections, tables, and body images in document order."""

    try:
        from docx import Document as WordDocument
        from docx.document import Document as DocumentProxy
        from docx.table import Table as WordTable
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn
    except ImportError as exc:  # pragma: no cover - packaging failure
        raise RuntimeError("python-docx is required for DOCX extraction") from exc

    source = Path(path)
    try:
        word_document = WordDocument(source)
    except Exception as exc:
        raise ValueError(f"could not parse DOCX {source}: {exc}") from exc

    matcher = HeadingMatcher(known_headings)
    result = Document(
        source_path=str(source),
        format="docx",
        metadata={
            "paragraph_count": len(word_document.paragraphs),
            "table_count": len(word_document.tables),
            "inline_shape_count": len(word_document.inline_shapes),
            "heading_matching": "style, outline/numbering metadata, configured exact normalized title, conservative numeric fallback",
        },
    )
    stack: list[Section] = []
    paragraph_index = 0
    table_index = 0

    def current_section() -> Section | None:
        return stack[-1] if stack else None

    def add_block(block: Block) -> None:
        result.raw_blocks.append(block)
        section = current_section()
        if section is not None:
            section.blocks.append(block)

    for child in word_document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, word_document)
            text = paragraph.text or ""
            style_name = getattr(paragraph.style, "name", None) if paragraph.style is not None else None
            level, recognition = _heading_level(paragraph, text, style_name, matcher, qn)
            numbering = extract_numbering(text)
            location = SourceLocation(kind="docx_paragraph", paragraph_index=paragraph_index)
            if level is not None:
                while stack and stack[-1].level >= level:
                    stack.pop()
                parent = current_section()
                parent_path = parent.path if parent else None
                section_path = f"{parent_path} / {text.strip()}" if parent_path else text.strip()
                section = Section(
                    title=text.strip(),
                    normalized_title=normalize_heading(text),
                    numbering=numbering,
                    level=level,
                    source_location=location,
                    parent_path=parent_path,
                    path=section_path,
                    metadata={"heading_recognition": recognition},
                )
                if parent is None:
                    result.sections.append(section)
                else:
                    parent.children.append(section)
                stack.append(section)
                block = Block(
                    kind="heading",
                    text=text,
                    source_location=location,
                    style_name=style_name,
                    level=level,
                    numbering=numbering,
                    section_path=section.path,
                    metadata={"heading_recognition": recognition},
                )
                add_block(block)
            else:
                section = current_section()
                block = Block(
                    kind="paragraph",
                    text=text,
                    source_location=location,
                    style_name=style_name,
                    numbering=numbering,
                    section_path=section.path if section else None,
                )
                add_block(block)
            paragraph_images = _images_for_element(
                paragraph._p,
                word_document,
                location,
                current_section().path if current_section() else None,
                qn,
            )
            for image in paragraph_images:
                result.images.append(image)
                add_block(
                    Block(
                        kind="image",
                        text=image.description or "",
                        source_location=image.source_location,
                        section_path=image.parent_section,
                        metadata=image.to_dict(),
                    )
                )
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            word_table = WordTable(child, word_document)
            section = current_section()
            location = SourceLocation(kind="docx_table", table_index=table_index)
            table = _extract_table(word_table, table_index, location, section.path if section else None, qn)
            result.tables.append(table)
            add_block(
                Block(
                    kind="table",
                    text=_table_text(table),
                    source_location=location,
                    section_path=section.path if section else None,
                    metadata={"dimensions": table.dimensions},
                )
            )
            for row_number, row in enumerate(word_table.rows, start=1):
                for column_number, cell in enumerate(row.cells, start=1):
                    cell_images = _images_for_element(
                        cell._tc,
                        word_document,
                        SourceLocation(
                            kind="docx_table_cell",
                            table_index=table_index,
                            row=row_number,
                            column=column_number,
                        ),
                        section.path if section else None,
                        qn,
                    )
                    for image in cell_images:
                        result.images.append(image)
                        add_block(
                            Block(
                                kind="image",
                                text=image.description or "",
                                source_location=image.source_location,
                                section_path=image.parent_section,
                                metadata=image.to_dict(),
                            )
                        )
            table_index += 1

    result.metadata["section_count"] = sum(1 for _ in result.all_sections())
    if classifier is not None:
        apply_classifier(result, classifier)
    return result


def apply_classifier(document: Document, classifier: ContentClassifier) -> None:
    for block in document.raw_blocks:
        classifier.classify_block(block)
    for table in document.tables:
        classifier.classify_table(table)


def _heading_level(paragraph: Any, text: str, style_name: str | None, matcher: HeadingMatcher, qn: Any) -> tuple[int | None, str]:
    if not text.strip():
        return None, "empty_paragraph"
    if style_name:
        match = re.search(r"heading\s*([1-9])", style_name, re.IGNORECASE)
        if match:
            return int(match.group(1)), "heading_style"
        if style_name.casefold() in {"title", "subtitle"}:
            return 1, "title_style"

    ppr = paragraph._p.pPr
    if ppr is not None:
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is not None and outline.get(qn("w:val")) is not None:
            try:
                return int(outline.get(qn("w:val"))) + 1, "outline_level"
            except ValueError:
                pass
        # A numbered paragraph is only promoted when its style/outline already
        # indicates heading-like structure.  Plain list numbering is not enough.
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is not None and style_name and "heading" in style_name.casefold():
            ilvl = num_pr.find(qn("w:ilvl"))
            if ilvl is not None and ilvl.get(qn("w:val")) is not None:
                return int(ilvl.get(qn("w:val"))) + 1, "heading_numbering"

    configured = matcher.match(text)
    if configured is not None:
        return configured, "configured_normalized_heading" if normalize_heading(text) in matcher.known else "conservative_numbered_heading"
    return None, "paragraph"


def _extract_table(word_table: Any, table_index: int, location: SourceLocation, parent_section: str | None, qn: Any) -> Table:
    rows = []
    merged_regions: list[str] = []
    for row_number, row in enumerate(word_table.rows, start=1):
        cells = []
        for column_number, cell in enumerate(row.cells, start=1):
            cell_location = SourceLocation(
                kind="docx_table_cell",
                table_index=table_index,
                row=row_number,
                column=column_number,
            )
            from ..model import Cell

            cells.append(
                Cell(
                    value=cell.text,
                    row=row_number,
                    column=column_number,
                    source_location=cell_location,
                    original_text=cell.text,
                )
            )
            tc_pr = cell._tc.tcPr
            grid_span = tc_pr.find(qn("w:gridSpan")) if tc_pr is not None else None
            if grid_span is not None:
                value = grid_span.get(qn("w:val"))
                if value and value != "1":
                    merged_regions.append(f"row {row_number}, col {column_number}, span {value}")
        rows.append(cells)
    return Table(rows=rows, source_location=location, merged_regions=merged_regions, parent_section=parent_section)


def _table_text(table: Table) -> str:
    return "\n".join(" | ".join(cell.original_text for cell in row) for row in table.rows)


def _images_for_element(
    element: Any,
    word_document: Any,
    location: SourceLocation,
    parent_section: str | None,
    qn: Any,
) -> list[Image]:
    images: list[Image] = []
    for drawing in element.iter(qn("w:drawing")):
        rel_ids = [blip.get(qn("r:embed")) for blip in drawing.iter(qn("a:blip"))]
        doc_pr = next(iter(drawing.iter(qn("wp:docPr"))), None)
        description = None
        if doc_pr is not None:
            description = doc_pr.get("descr") or doc_pr.get("name")
        extent = next(iter(drawing.iter(qn("wp:extent"))), None)
        width = height = None
        if extent is not None:
            try:
                width = int(extent.get("cx"))
                height = int(extent.get("cy"))
            except (TypeError, ValueError):
                pass
        for relationship_id in rel_ids:
            related = word_document.part.related_parts.get(relationship_id)
            filename = None
            if related is not None:
                filename = str(getattr(related, "partname", "")) or None
            images.append(
                Image(
                    relationship_id=relationship_id,
                    description=description,
                    source_location=location,
                    parent_section=parent_section,
                    filename=filename,
                    width=width,
                    height=height,
                    metadata={"embedded": True},
                )
            )
    return images
